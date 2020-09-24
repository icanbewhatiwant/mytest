from docx import Document
import json
import re


def data2Excel(drug_dict,drug_list,pra_len):
    pra_lens = pra_len-1
    for i,par in enumerate(paragraphs): #获得段落对象列表
        par_str = par.text
        # if i <=50:
        if par_str =="" :
            continue

        #提取药品中、英文名称
        en_match = enpattern.search(par_str)
        if en_match:
            #匹配到药品英文名称时，字典内容作为一个整体存入list
            if drug_dict:
                drug_list.append(drug_dict)
                drug_dict={}

            en_name = par_str
            k = i + 1
            while True:#英文名称可能会换行
                if k <= pra_lens:
                    next_en_text = paragraphs[k].text
                    next_en_match = enpattern.search(next_en_text)
                    if next_en_match:
                        en_name += " /"+ next_en_text
                        k +=1
                    else:
                        break
            drug_dict["enName"] = en_name
            drug_dict["drugName"] = paragraphs[i-1].text
            # print("drug_name", drug_name)
            # print("en_name:", en_name)

        #提取各种标签及其对应内容
        match = label_patr.search(par_str)
        if match:
            label_str = match.group()
            #因为括号中可能有.或者其他符号，找到括号中的全部中文，拼接为label
            label_cmatch = label_con.findall(label_str)
            if label_cmatch:
                label = ''.join(label_cmatch)
            label_len = len(label_str)
            content = par_str[label_len:]

            #下一段可能还是继续当前标签内容，判断并拼接
            j = i+1
            while True:
                if j <=pra_lens:
                    next_text = paragraphs[j].text
                    # 下一段是否为空
                    if next_text =="":
                        j+=1
                        continue
                    # 匹配到第几节、第几章
                    if chapter_patr.search(next_text):
                        break
                    # 下下段匹配到药品英文名称，则下段是药品名称
                    if j+1 <=pra_lens:
                        nex_nex_text = paragraphs[j+1].text
                        nexen_match = enpattern.search(nex_nex_text)
                        if nexen_match:
                            break
                    # 下一段匹配到label
                    if label_patr.search(next_text):
                        break

                    content += "&nsp"+next_text
                    j+=1
                else:
                    break
            drug_dict[label] = content

            # print("label",label)
            # print("content",content)
            # print()
        # else:
        #     break
    if drug_list:
        with open("C:/产品文档/转换器测试数据/1401-1539补充数据.json","w",encoding='utf-8') as fp:
            fp.write(json.dumps(drug_list, indent=4,ensure_ascii=False))#unicode串转中文传入

#遍历文件得到label名称，去重作为列名

# 得到：['【药理】', '【不良反应】', '【禁忌证】', '【注意事项】', '【用法与用量】', '【儿科用法与用量】', '【制剂与规格】', '【适应证】', '【儿科注意事项】', '【药物相互作用】', '【给药说明】']
def getLabellist():
    label_list=[]

    for i, par in enumerate(paragraphs):
        par_str = par.text
        if par_str == "":
            continue

            # 提取各种标签及其对应内容
        match = label_patr.search(par_str)
        if match:
            label_str = match.group()
            label_cmatch = label_con.findall(label_str)
            # if label_cmatch:
            #     label = ''.join(label_cmatch)
            #     if label not in label_list:
            #         label_list.append(label)
            if label_str not in label_list:
                label_list.append(label_str)

    print(label_list)


if __name__ == "__main__":
    # Document 类，不仅可以新建word文档，也可以打开一个本地文档
    # doc = Document('C:/产品文档/转换器测试数据/2015年版_1401-1539.docx')
    doc = Document('C:/产品文档/转换器测试数据/2015年版_1201-1400.docx')

    "获取文档所有段落信息："
    # 获取文档所有段落对象
    paragraphs = doc.paragraphs
    pra_len = len(paragraphs)

    # 源数据是pdf转word，数据不规范
    # 标签，前面括号格式有中英文全半角的各种格式组合
    # label_patr = re.compile("^(【|\[|\[:|［:|［|.【)[\u4e00-\u9fa5]+(】|\]|］)")#完全匹配中间为中文
    # label_patr = re.compile("^(【|\[|\[:|［:|［|.【).+(】|\]|］)")#匹配中间为任意字符，文字中有】会匹配到文字中，贪婪
    label_patr = re.compile("^(.{,2}【|\[|\[:|［:|［|.【)[^】|\]|］]+(】|\]|］)")  # 中间不能有】，实现非贪婪
    label_con = re.compile("[\u4e00-\u9fa5]+")
    # 匹配药品英文名称
    enpattern = re.compile(u'^[a-zA-Z\s/-]+$')#包含"/"

    # 匹配章、节
    chapter_patr = re.compile('第.+章|第.+节')

    drug_dict = {}
    drug_list = []

    getLabellist()
    # data2Excel(drug_dict,drug_list,pra_len)




